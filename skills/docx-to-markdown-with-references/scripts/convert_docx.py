#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
python-docx を使った Word → Markdown 変換（pandoc がない環境用フォールバック）

使用方法:
    python convert_docx.py <input.docx> <output.md>

依存関係: pip install python-docx
"""

import sys
from docx import Document


def convert_docx_to_markdown(input_path: str, output_path: str) -> None:
    doc = Document(input_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if "Heading 1" in style_name or style_name == "見出し 1":
            lines.append(f"# {text}")
        elif "Heading 2" in style_name or style_name == "見出し 2":
            lines.append(f"## {text}")
        elif "Heading 3" in style_name or style_name == "見出し 3":
            lines.append(f"### {text}")
        elif "Heading 4" in style_name or style_name == "見出し 4":
            lines.append(f"#### {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        table_lines = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            table_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                table_lines.append("|" + "|".join(["---" for _ in cells]) + "|")
        lines.append("\n" + "\n".join(table_lines) + "\n")

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <input.docx> <output.md>")
        sys.exit(1)

    convert_docx_to_markdown(sys.argv[1], sys.argv[2])
    print(f"変換完了: {sys.argv[2]}")
